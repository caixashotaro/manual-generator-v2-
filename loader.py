"""
loader.py - 多形式ナレッジファイル読み込みモジュール

対応形式:
  - CSV (.csv)
  - Excel (.xlsx, .xls)
  - PDF (.pdf)
  - Word (.docx)
  - テキスト (.txt, .md)

使い方:
  from loader import load_knowledge_folder
  knowledge_text = load_knowledge_folder("./knowledge")
"""

import os
import sys
from pathlib import Path

# CSV / Excel
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

# PDF
try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# Word
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ============================================
# 個別ファイル読み込み関数
# ============================================

def load_csv(filepath: str) -> str:
    """CSVファイルを読み込み、テキスト化する"""
    if not PANDAS_AVAILABLE:
        raise ImportError("pandas が必要です: pip install pandas")
    try:
        df = pd.read_csv(filepath, encoding="utf-8")
    except UnicodeDecodeError:
        df = pd.read_csv(filepath, encoding="cp932")

    return _dataframe_to_text(df)


def load_excel(filepath: str) -> str:
    """Excelファイル (.xlsx/.xls) を読み込み、全シートをテキスト化する"""
    if not PANDAS_AVAILABLE:
        raise ImportError("pandas が必要です: pip install pandas openpyxl")

    sheets = pd.read_excel(filepath, sheet_name=None, engine="openpyxl")
    parts = []
    for sheet_name, df in sheets.items():
        parts.append(f"--- シート: {sheet_name} ---")
        parts.append(_dataframe_to_text(df))
    return "\n".join(parts)


def load_pdf(filepath: str) -> str:
    """PDFファイルからテキストを抽出する"""
    if not PYPDF2_AVAILABLE:
        raise ImportError("PyPDF2 が必要です: pip install PyPDF2")
    reader = PdfReader(filepath)
    pages = []
    for i, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text and text.strip():
            pages.append(f"[ページ {i}]\n{text.strip()}")
    return "\n\n".join(pages)


def load_docx(filepath: str) -> str:
    """Wordファイル (.docx) からテキストを抽出する"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx が必要です: pip install python-docx")
    doc = Document(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())

    # テーブルも読み込む
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            paragraphs.append("\n".join(rows))

    return "\n".join(paragraphs)


def load_text(filepath: str) -> str:
    """テキストファイル (.txt, .md) を読み込む"""
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read().strip()
    except UnicodeDecodeError:
        with open(filepath, "r", encoding="cp932") as f:
            return f.read().strip()


# ============================================
# ユーティリティ
# ============================================

def _dataframe_to_text(df: "pd.DataFrame") -> str:
    """DataFrameをテキスト表現に変換する

    ヘッダー行 + 各データ行を読みやすい形式で出力。
    """
    if df.empty:
        return "(空のデータ)"

    # NaN を空文字に置換
    df = df.fillna("")

    lines = []
    # カラム名
    columns = list(df.columns)
    lines.append("列名: " + " | ".join(str(c) for c in columns))
    lines.append("-" * 40)

    # データ行（最大500行まで）
    max_rows = 500
    for idx, row in df.head(max_rows).iterrows():
        row_parts = []
        for col in columns:
            val = str(row[col]).strip()
            if val:
                row_parts.append(f"{col}: {val}")
        if row_parts:
            lines.append(f"行{idx + 1}: " + " / ".join(row_parts))

    if len(df) > max_rows:
        lines.append(f"... (以降 {len(df) - max_rows} 行省略)")

    return "\n".join(lines)


# ============================================
# フォルダ一括読み込み
# ============================================

# 拡張子 → 読み込み関数のマッピング
LOADERS = {
    ".csv": load_csv,
    ".xlsx": load_excel,
    ".xls": load_excel,
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".txt": load_text,
    ".md": load_text,
}


def load_single_file(filepath: str) -> str:
    """単一ファイルを読み込んでテキスト化する"""
    ext = Path(filepath).suffix.lower()
    loader = LOADERS.get(ext)
    if loader is None:
        print(f"  [スキップ] 非対応形式: {filepath}")
        return ""
    try:
        content = loader(filepath)
        return content
    except Exception as e:
        print(f"  [エラー] {filepath}: {e}")
        return ""


def load_knowledge_folder(folder_path: str) -> str:
    """フォルダ内の全対応ファイルを読み込み、統合テキストを返す

    Args:
        folder_path: ナレッジファイルが格納されたフォルダパス

    Returns:
        「ファイル名: 内容」の形式で結合されたテキスト
    """
    folder = Path(folder_path)
    if not folder.is_dir():
        raise FileNotFoundError(f"フォルダが見つかりません: {folder_path}")

    parts = []
    file_count = 0
    supported_exts = set(LOADERS.keys())

    # フォルダを再帰的に探索
    for root, dirs, files in os.walk(folder_path):
        # 隠しフォルダをスキップ
        dirs[:] = [d for d in dirs if not d.startswith(".")]
        for filename in sorted(files):
            if filename.startswith("."):
                continue
            ext = Path(filename).suffix.lower()
            if ext not in supported_exts:
                continue

            filepath = os.path.join(root, filename)
            rel_path = os.path.relpath(filepath, folder_path)
            print(f"  読み込み中: {rel_path}")

            content = load_single_file(filepath)
            if content:
                parts.append(f"========== ファイル: {rel_path} ==========\n{content}")
                file_count += 1

    if file_count == 0:
        print(f"  警告: {folder_path} に対応ファイルが見つかりませんでした。")
        return ""

    print(f"  合計 {file_count} ファイルを読み込みました。")
    combined = "\n\n".join(parts)
    print(f"  テキスト総量: {len(combined):,} 文字")
    return combined


# ============================================
# CLI として単独実行
# ============================================

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="ナレッジフォルダを読み込んでテキスト化")
    parser.add_argument("folder", help="ナレッジファイルが格納されたフォルダパス")
    parser.add_argument("-o", "--output", help="出力テキストファイルパス（省略時は標準出力）")
    args = parser.parse_args()

    print(f"ナレッジ読み込み開始: {args.folder}")
    text = load_knowledge_folder(args.folder)

    if args.output:
        with open(args.output, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"出力完了: {args.output} ({len(text):,} 文字)")
    else:
        print("\n" + "=" * 60)
        print(text)
