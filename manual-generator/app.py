"""
業務マニュアル自動生成アプリ
動画（MP4）または音声（MP3）から判断基準が含まれる高度な業務マニュアルを自動生成

アーキテクチャ:
- 音声認識: Groq Whisper API (large-v3) - 高精度・高速
- テキスト分析: Google Gemini API
- ホスティング: Streamlit Cloud対応

Requirements (requirements.txt):
streamlit>=1.28.0
opencv-python-headless>=4.8.0
groq>=0.4.0
google-generativeai>=0.3.0
Pillow>=10.0.0
python-docx>=1.0.0
imageio-ffmpeg>=0.4.9
"""

import streamlit as st
import cv2
import tempfile
import os
import json
import base64
import zipfile
import io
import subprocess
import re
import shutil
import glob as glob_module
from datetime import datetime, timedelta
from PIL import Image
import numpy as np

# Word出力用
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ffmpegのパスを取得してPATHに追加
try:
    import imageio_ffmpeg
    FFMPEG_PATH = imageio_ffmpeg.get_ffmpeg_exe()
    ffmpeg_dir = os.path.dirname(FFMPEG_PATH)
    os.environ["PATH"] = ffmpeg_dir + os.pathsep + os.environ.get("PATH", "")
except ImportError:
    FFMPEG_PATH = "ffmpeg"

# Groq API (Whisper)
try:
    from groq import Groq
    GROQ_AVAILABLE = True
except ImportError:
    GROQ_AVAILABLE = False

# Google GenAI SDK (Gemini) - 新しい公式SDK
try:
    from google import genai
    from google.genai import types
    GENAI_AVAILABLE = True
except ImportError:
    GENAI_AVAILABLE = False

# Supabase (クラウドDB)
try:
    from supabase import create_client, Client
    SUPABASE_AVAILABLE = True
except ImportError:
    SUPABASE_AVAILABLE = False


# ファイルサイズ制限（Groq API: 25MB）
MAX_CHUNK_SIZE_MB = 25
MAX_CHUNK_SIZE_BYTES = MAX_CHUNK_SIZE_MB * 1024 * 1024

# Gemini モデル設定
# 最高精度モデル: gemini-2.5-pro（2025年最新）
GEMINI_MODEL_PRO = "gemini-2.5-pro"

# ============================================
# Supabase クライアント管理
# ============================================

def get_supabase_client() -> Client:
    """Supabaseクライアントを取得"""
    if not SUPABASE_AVAILABLE:
        return None

    # Streamlit secretsまたはsession_stateからURLとキーを取得
    supabase_url = None
    supabase_key = None

    # secretsから読み込み（エラーを無視）
    try:
        supabase_url = st.secrets.get("SUPABASE_URL", "")
        supabase_key = st.secrets.get("SUPABASE_KEY", "")
    except Exception:
        pass

    # session_stateから読み込み（secretsより優先）
    if not supabase_url and "supabase_url" in st.session_state:
        supabase_url = st.session_state.supabase_url
    if not supabase_key and "supabase_key" in st.session_state:
        supabase_key = st.session_state.supabase_key

    if supabase_url and supabase_key:
        try:
            return create_client(supabase_url, supabase_key)
        except Exception:
            return None
    return None


# ============================================
# ローカルプロジェクトストレージ
# ============================================

PROJECTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "projects")
os.makedirs(PROJECTS_DIR, exist_ok=True)


def _project_json_path(project_id: str) -> str:
    """プロジェクトJSONファイルのパスを返す"""
    return os.path.join(PROJECTS_DIR, f"{project_id}.json")


def _project_source_path(project_id: str, ext: str = ".mp4") -> str:
    """プロジェクト動画/音声ファイルのパスを返す"""
    return os.path.join(PROJECTS_DIR, f"{project_id}_source{ext}")


def get_local_project_list() -> list:
    """projects/ 内のJSONファイル一覧を読み込んでプロジェクトリストを返す"""
    projects = []
    for json_path in glob_module.glob(os.path.join(PROJECTS_DIR, "*.json")):
        project_id = os.path.splitext(os.path.basename(json_path))[0]
        try:
            with open(json_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            projects.append({
                "id": project_id,
                "name": data.get("name", "無題のプロジェクト"),
                "created_at": data.get("created_at", ""),
                "video_name": data.get("video_name", ""),
            })
        except Exception:
            pass
    # created_at の降順でソート
    projects.sort(key=lambda p: p.get("created_at", ""), reverse=True)
    return projects


def save_project_local(project_id: str, name: str, video_name: str,
                       steps: list, flow_summary: str, segments: list,
                       video_duration: float, video_path: str = None,
                       is_audio_only: bool = False):
    """プロジェクトデータをローカルJSONに保存し、動画/音声をprojects/にコピーする"""
    # 画像をBase64エンコード
    steps_to_save = []
    for step in steps:
        step_copy = step.copy()
        if step_copy.get("image") is not None:
            try:
                pil_image = Image.fromarray(step_copy["image"])
                buf = io.BytesIO()
                pil_image.save(buf, format="PNG", optimize=True)
                step_copy["image_base64"] = base64.b64encode(buf.getvalue()).decode()
            except Exception:
                step_copy["image_base64"] = step_copy.get("image_base64")
            del step_copy["image"]
        else:
            if "image_base64" not in step_copy:
                step_copy["image_base64"] = None
            if "image" in step_copy:
                del step_copy["image"]
        steps_to_save.append(step_copy)

    # 既存JSONがあればcreated_atを引き継ぐ
    json_path = _project_json_path(project_id)
    created_at = datetime.now().isoformat()
    if os.path.exists(json_path):
        try:
            with open(json_path, "r", encoding="utf-8") as f:
                old = json.load(f)
            created_at = old.get("created_at", created_at)
        except Exception:
            pass

    project_data = {
        "name": name,
        "video_name": video_name,
        "created_at": created_at,
        "flow_summary": flow_summary,
        "steps": steps_to_save,
        "segments": segments,
        "video_duration": video_duration,
    }

    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(project_data, f, ensure_ascii=False)

    # 動画/音声ファイルをprojects/にコピー
    if video_path and os.path.exists(video_path):
        ext = ".mp3" if is_audio_only else ".mp4"
        dest = _project_source_path(project_id, ext)
        if os.path.abspath(video_path) != os.path.abspath(dest):
            shutil.copy2(video_path, dest)

    return True


def load_project_local(project_id: str) -> dict:
    """ローカルJSONからプロジェクトを読み込み、画像を復元する"""
    json_path = _project_json_path(project_id)
    if not os.path.exists(json_path):
        return None

    try:
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        return None

    # Base64画像をnumpy配列に復元
    for step in data.get("steps", []):
        if step.get("image_base64"):
            try:
                img_data = base64.b64decode(step["image_base64"])
                pil_image = Image.open(io.BytesIO(img_data))
                step["image"] = np.array(pil_image)
            except Exception:
                step["image"] = None
        else:
            step["image"] = None
        if "image_base64" in step:
            del step["image_base64"]

    # 動画/音声ファイルパスを復元
    for ext in [".mp4", ".mp3"]:
        source = _project_source_path(project_id, ext)
        if os.path.exists(source):
            data["video_path"] = source
            data["is_audio_only"] = (ext == ".mp3")
            break

    data["id"] = project_id
    return data


def delete_project_local(project_id: str) -> bool:
    """ローカルのJSON + 動画/音声ファイルを削除"""
    json_path = _project_json_path(project_id)
    deleted = False
    if os.path.exists(json_path):
        os.remove(json_path)
        deleted = True
    for ext in [".mp4", ".mp3"]:
        source = _project_source_path(project_id, ext)
        if os.path.exists(source):
            os.remove(source)
    return deleted


# ============================================
# プロジェクト管理機能（ローカル優先 + Supabase連携）
# ============================================

def get_project_list() -> list:
    """ローカルプロジェクト一覧を返す（Supabaseがあれば統合）"""
    projects = get_local_project_list()
    local_ids = {p["id"] for p in projects}

    supabase = get_supabase_client()
    if supabase:
        try:
            response = supabase.table("projects").select(
                "id, name, created_at, video_name"
            ).order("created_at", desc=True).execute()
            for row in response.data:
                if row["id"] not in local_ids:
                    projects.append({
                        "id": row["id"],
                        "name": row.get("name", "無題のプロジェクト"),
                        "created_at": row.get("created_at", ""),
                        "video_name": row.get("video_name", ""),
                    })
        except Exception:
            pass

    projects.sort(key=lambda p: p.get("created_at", ""), reverse=True)
    return projects


def save_project(project_id: str, name: str, video_name: str, steps: list,
                 flow_summary: str, segments: list, video_duration: float):
    """常にローカル保存 + Supabaseがあればクラウドにも保存"""
    # ローカル保存
    save_project_local(
        project_id=project_id,
        name=name,
        video_name=video_name,
        steps=steps,
        flow_summary=flow_summary,
        segments=segments,
        video_duration=video_duration,
        video_path=st.session_state.get("video_path"),
        is_audio_only=st.session_state.get("is_audio_only", False),
    )

    # Supabase保存（任意）
    supabase = get_supabase_client()
    if supabase:
        steps_to_save = []
        for step in steps:
            step_copy = step.copy()
            if step_copy.get("image") is not None:
                try:
                    pil_image = Image.fromarray(step_copy["image"])
                    buf = io.BytesIO()
                    pil_image.save(buf, format="PNG", optimize=True)
                    step_copy["image_base64"] = base64.b64encode(buf.getvalue()).decode()
                except Exception:
                    step_copy["image_base64"] = None
                del step_copy["image"]
            else:
                step_copy["image_base64"] = None
                if "image" in step_copy:
                    del step_copy["image"]
            steps_to_save.append(step_copy)

        project_data = {
            "id": project_id,
            "name": name,
            "video_name": video_name,
            "flow_summary": flow_summary,
            "steps": steps_to_save,
            "segments": segments,
            "video_duration": video_duration,
        }
        try:
            supabase.table("projects").upsert(project_data).execute()
        except Exception:
            pass

    return True


def load_project(project_id: str) -> dict:
    """ローカルから読み込み（なければSupabase）"""
    data = load_project_local(project_id)
    if data:
        return data

    supabase = get_supabase_client()
    if not supabase:
        return None

    try:
        response = supabase.table("projects").select("*").eq("id", project_id).execute()
        if not response.data:
            return None

        data = response.data[0]
        for step in data.get("steps", []):
            if step.get("image_base64"):
                try:
                    img_data = base64.b64decode(step["image_base64"])
                    pil_image = Image.open(io.BytesIO(img_data))
                    step["image"] = np.array(pil_image)
                except Exception:
                    step["image"] = None
            else:
                step["image"] = None
            if "image_base64" in step:
                del step["image_base64"]

        data["steps"] = data.get("steps", [])
        return data
    except Exception:
        return None


def delete_project(project_id: str) -> bool:
    """ローカル削除 + Supabase削除"""
    deleted = delete_project_local(project_id)

    supabase = get_supabase_client()
    if supabase:
        try:
            supabase.table("projects").delete().eq("id", project_id).execute()
            deleted = True
        except Exception:
            pass

    return deleted


def generate_project_id() -> str:
    """ユニークなプロジェクトIDを生成"""
    import uuid
    return str(uuid.uuid4())[:8]


def format_timestamp(seconds: float) -> str:
    """秒数をMM:SS形式に変換"""
    td = timedelta(seconds=int(seconds))
    minutes, secs = divmod(int(seconds), 60)
    return f"{minutes:02d}:{secs:02d}"


def get_file_size(file_path: str) -> int:
    """ファイルサイズをバイトで取得"""
    return os.path.getsize(file_path)


def extract_audio_mp3(video_path: str, output_path: str) -> bool:
    """動画全体から音声をmp3で一括抽出（Groq API対応、低ビットレートで圧縮）"""
    try:
        cmd = [
            FFMPEG_PATH, "-y",
            "-i", video_path,
            "-vn",
            "-acodec", "libmp3lame",
            "-ar", "16000",
            "-ac", "1",
            "-b:a", "48k",
            output_path
        ]
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode == 0 and os.path.exists(output_path):
            return True
        else:
            st.error(f"音声抽出エラー: {result.stderr}")
            return False
    except Exception as e:
        st.error(f"音声抽出エラー: {str(e)}")
        return False


def split_audio_file(audio_path: str, max_size_bytes: int = MAX_CHUNK_SIZE_BYTES) -> list:
    """
    音声ファイルが25MBを超える場合のみ時間で分割
    Returns: [(chunk_path, start_offset, is_temp), ...]
    """
    file_size = get_file_size(audio_path)

    if file_size <= max_size_bytes:
        return [(audio_path, 0, False)]

    # 音声の長さを取得
    cmd = [FFMPEG_PATH, "-i", audio_path, "-f", "null", "-"]
    result = subprocess.run(cmd, capture_output=True, text=True)
    # ffmpegの出力からdurationを解析
    duration = 0
    for line in result.stderr.split('\n'):
        if 'Duration' in line:
            time_str = line.split('Duration:')[1].split(',')[0].strip()
            parts = time_str.split(':')
            duration = float(parts[0]) * 3600 + float(parts[1]) * 60 + float(parts[2])
            break

    if duration == 0:
        return [(audio_path, 0, False)]

    # 分割数を計算
    num_chunks = int(np.ceil(file_size / max_size_bytes))
    chunk_duration = duration / num_chunks

    chunks = []
    temp_dir = tempfile.mkdtemp()

    for i in range(num_chunks):
        start_time = i * chunk_duration
        chunk_length = min(chunk_duration, duration - start_time)
        chunk_path = os.path.join(temp_dir, f"audio_chunk_{i}.mp3")

        cmd = [
            FFMPEG_PATH, "-y",
            "-ss", str(start_time),
            "-i", audio_path,
            "-t", str(chunk_length),
            "-acodec", "copy",
            chunk_path
        ]
        result = subprocess.run(cmd, capture_output=True, text=True)

        if result.returncode == 0 and os.path.exists(chunk_path):
            chunks.append((chunk_path, start_time, True))
        else:
            st.warning(f"音声チャンク {i+1} の作成に失敗しました")

    return chunks if chunks else [(audio_path, 0, False)]


def transcribe_full_video(video_path: str, groq_api_key: str,
                          whisper_model: str = "whisper-large-v3",
                          progress_callback=None, status_callback=None) -> list:
    """
    動画全体から音声を一括抽出し、Groq Whisper APIで文字起こし
    動画ファイルのサイズに関係なく、音声レベルで最適な分割を行う
    """
    # Step 1: 動画全体から音声をmp3で一括抽出
    if status_callback:
        status_callback("動画から音声を抽出中...")
    if progress_callback:
        progress_callback(10)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp_audio:
        audio_path = tmp_audio.name

    if not extract_audio_mp3(video_path, audio_path):
        return []

    audio_size_mb = get_file_size(audio_path) / (1024 * 1024)
    if status_callback:
        status_callback(f"音声抽出完了（{audio_size_mb:.1f}MB）")

    # Step 2: 音声ファイルが25MB超の場合のみ分割
    audio_chunks = split_audio_file(audio_path)
    num_chunks = len(audio_chunks)

    if num_chunks > 1 and status_callback:
        status_callback(f"音声を {num_chunks} 個に分割して処理します...")

    # Step 3: 各音声チャンクをGroq APIで文字起こし
    all_segments = []
    for chunk_idx, (chunk_path, start_offset, is_temp) in enumerate(audio_chunks):
        if status_callback:
            if num_chunks > 1:
                status_callback(f"Groq Whisper APIで音声認識中... ({chunk_idx + 1}/{num_chunks})")
            else:
                status_callback("Groq Whisper APIで音声認識中...")

        if progress_callback:
            progress_callback(15 + int((chunk_idx / num_chunks) * 25))

        segments = transcribe_audio_with_groq(chunk_path, groq_api_key, whisper_model)

        # タイムスタンプをオフセット調整
        for segment in segments:
            segment["start"] += start_offset
            segment["end"] += start_offset
            all_segments.append(segment)

        # 一時ファイルを削除
        if is_temp:
            try:
                os.unlink(chunk_path)
            except:
                pass

    # 元の音声ファイルを削除
    try:
        os.unlink(audio_path)
    except:
        pass

    return all_segments


def transcribe_audio_with_groq(audio_path: str, groq_api_key: str, model: str = "whisper-large-v3") -> list:
    """Groq Whisper APIで音声を文字起こし（タイムスタンプ付き）"""
    if not GROQ_AVAILABLE:
        st.error("Groq SDKがインストールされていません。`pip install groq`を実行してください。")
        return []

    try:
        client = Groq(api_key=groq_api_key)

        with open(audio_path, "rb") as audio_file:
            # verbose_jsonでタイムスタンプ付きの結果を取得
            transcription = client.audio.transcriptions.create(
                file=audio_file,
                model=model,
                response_format="verbose_json",
                language="ja",
                temperature=0.0
            )

        segments = []

        # verbose_jsonの場合、segmentsが含まれる
        if hasattr(transcription, 'segments') and transcription.segments:
            for segment in transcription.segments:
                segments.append({
                    "start": segment.get("start", 0),
                    "end": segment.get("end", 0),
                    "text": segment.get("text", "")
                })
        elif hasattr(transcription, 'text'):
            # セグメントがない場合は全体を1つのセグメントとして扱う
            segments.append({
                "start": 0,
                "end": 0,
                "text": transcription.text
            })

        return segments
    except Exception as e:
        st.error(f"Groq音声認識エラー: {str(e)}")
        return []


def transcribe_audio_file(audio_path: str, groq_api_key: str,
                          whisper_model: str = "whisper-large-v3",
                          progress_callback=None, status_callback=None) -> list:
    """
    音声ファイル（MP3）を直接Groq Whisper APIで文字起こし
    必要に応じて最適化（16kHz モノラル化）を行う
    """
    if status_callback:
        status_callback("音声ファイルを最適化中...")
    if progress_callback:
        progress_callback(10)

    # 音声を最適化（16kHz モノラル、低ビットレート）してサイズを削減
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp_audio:
        optimized_path = tmp_audio.name

    try:
        cmd = [
            FFMPEG_PATH, "-y",
            "-i", audio_path,
            "-vn",
            "-acodec", "libmp3lame",
            "-ar", "16000",
            "-ac", "1",
            "-b:a", "48k",
            optimized_path
        ]
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0 or not os.path.exists(optimized_path):
            optimized_path = audio_path
    except Exception:
        optimized_path = audio_path

    audio_size_mb = get_file_size(optimized_path) / (1024 * 1024)
    if status_callback:
        status_callback(f"音声ファイル（{audio_size_mb:.1f}MB）を処理します")

    # 音声ファイルが25MB超の場合のみ分割
    audio_chunks = split_audio_file(optimized_path)
    num_chunks = len(audio_chunks)

    if num_chunks > 1 and status_callback:
        status_callback(f"音声を {num_chunks} 個に分割して処理します...")

    # 各音声チャンクをGroq APIで文字起こし
    all_segments = []
    for chunk_idx, (chunk_path, start_offset, is_temp) in enumerate(audio_chunks):
        if status_callback:
            if num_chunks > 1:
                status_callback(f"Groq Whisper APIで音声認識中... ({chunk_idx + 1}/{num_chunks})")
            else:
                status_callback("Groq Whisper APIで音声認識中...")

        if progress_callback:
            progress_callback(15 + int((chunk_idx / num_chunks) * 25))

        segments = transcribe_audio_with_groq(chunk_path, groq_api_key, whisper_model)

        for segment in segments:
            segment["start"] += start_offset
            segment["end"] += start_offset
            all_segments.append(segment)

        if is_temp:
            try:
                os.unlink(chunk_path)
            except:
                pass

    # 最適化された一時ファイルを削除
    if optimized_path != audio_path:
        try:
            os.unlink(optimized_path)
        except:
            pass

    return all_segments


def extract_frame(video_path: str, timestamp: float) -> np.ndarray:
    """指定タイムスタンプのフレームを抽出"""
    cap = cv2.VideoCapture(video_path)
    cap.set(cv2.CAP_PROP_POS_MSEC, timestamp * 1000)
    ret, frame = cap.read()
    cap.release()

    if ret:
        return cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
    return None


def get_video_duration(video_path: str) -> float:
    """動画の長さを取得"""
    cap = cv2.VideoCapture(video_path)
    fps = cap.get(cv2.CAP_PROP_FPS)
    frame_count = cap.get(cv2.CAP_PROP_FRAME_COUNT)
    cap.release()

    if fps > 0:
        return frame_count / fps
    return 0


def get_audio_duration(audio_path: str) -> float:
    """音声ファイルの長さを取得（ffmpegを使用）"""
    try:
        cmd = [FFMPEG_PATH, "-i", audio_path, "-f", "null", "-"]
        result = subprocess.run(cmd, capture_output=True, text=True)
        for line in result.stderr.split('\n'):
            if 'Duration' in line:
                time_str = line.split('Duration:')[1].split(',')[0].strip()
                parts = time_str.split(':')
                return float(parts[0]) * 3600 + float(parts[1]) * 60 + float(parts[2])
    except Exception:
        pass
    return 0


def build_full_transcript(segments: list) -> str:
    """
    Whisperのセグメントから、タイムスタンプ付きの全文テキストを生成
    フォーマット: [MM:SS-MM:SS] テキスト
    """
    lines = []
    for seg in segments:
        start_ts = format_timestamp(seg["start"])
        end_ts = format_timestamp(seg["end"])
        text = seg["text"].strip()
        if text:
            lines.append(f"[{start_ts}-{end_ts}] {text}")
    return "\n".join(lines)


def analyze_full_context_with_gemini(api_key: str, full_transcript: str) -> dict:
    """
    【Stage 1】フルコンテクストをGemini Proに渡して、業務フローの流れとアクション列を生成

    Returns:
        {
            "flow_summary": "業務全体の流れを数文で説明した要約",
            "actions": [
                {
                    "index": 1,
                    "action_title": "...",
                    "description": "...(完結した文章)...",
                    "actor": "担当者名や役割",
                    "screen": "画面名やシステム名",
                    "object": "操作対象物",
                    "operation": "具体的な操作内容",
                    "ai_hypothesis": "...",
                    "importance": "high" or "normal"
                },
                ...
            ]
        }

    ※ この段階ではタイムスタンプは付けない（テキストベースのアクション定義のみ）
    """
    if not GENAI_AVAILABLE:
        st.error("Google GenAI SDKがインストールされていません。`pip install google-genai`を実行してください。")
        return None

    try:
        client = genai.Client(api_key=api_key)

        prompt = f"""
あなたは業務マニュアル作成の専門家です。

以下は、ある業務操作を録画した動画から音声認識（Whisper）で取得した全文の文字起こしです。
タイムスタンプ付きの発話ログを読み、**動画全体のコンテクスト（流れ・ストーリー）** を把握したうえで、
業務フロー全体の要約と、そのフローに沿ったアクション（操作）の一覧を作成してください。

【重要な指示 - アクション分割の粒度】
1. まず動画全体を通して「この業務は何をしているのか」「どんな流れで進んでいるのか」を理解してください。
2. **アクションは可能な限り細かく分割してください。**
   - 1つのアクションの中で、「登場人物」「対象物」「操作内容」「目的」が1つに絞り込めるところまで細かく分割すること。
   - 1つの説明文に複数の操作が混ざってしまう場合は、操作ごとに別のアクションとして分割すること。
3. **各アクションには、以下の情報を可能な限り具体的に含めてください：**
   - 誰が（担当者、役職、部署など）
   - いつ（どのタイミングで、どの条件で）
   - どの画面で（システム名、画面名、タブ名など）
   - どの項目を（ボタン名、フィールド名、書類名など）
   - どう操作するか（クリック、入力、確認、分類など）

【重要な指示 - 入力・選択操作の詳細化】
**特定の画面でフォーム入力や選択操作を行っている場面では、以下のレベルまで詳細に説明してください：**
- 「どの画面」の
- 「どの項目（ラベル名・フィールド名）」に
- 「どんな値・情報」を
- 「なぜそのように入力／選択しているのか」

**入力・選択操作の説明の良い例：**
- 「`受付日` 欄に、申請書に記載されている受付日をカレンダーから選択して入力します。」
- 「`担当者` プルダウンから、自分の名前を選択します。」
- 「`申請番号` フィールドに、申請書右上に記載されている8桁の番号を入力します。」
- 「`検索` ボタンをクリックして、条件に合致する申請を一覧表示します。」

**入力・選択操作の悪い例（このレベルは避けること）：**
- 「この画面で条件を設定して検索します」（→ どの項目に何を入力するか不明）
- 「必要事項を入力します」（→ 具体的な項目名や値が不明）

**ただし、UIから明確に読み取れない項目名や値を、勝手に創作しないでください。**

【重要な指示 - 項目名・リストの完全列挙】
**顧客から直接指示されている項目名・分類名・チェック項目などが音声や画面から読み取れる場合は、それらを省略せず、すべて列挙してください。**

- 「など」「〜等」「その他」といったまとめ表現は使わず、実際の項目名をすべて書き出すこと。
- 例えば：
  - 必須チェック項目として「氏名」「住所」「電話番号」「生年月日」「申請日」が言及されている場合、5つすべてを列挙する。
  - FAX送信先として「総務課」「経理課」「営業部」が言及されている場合、3つすべてを列挙する。
- 箇条書きでも、1文の中にカンマ区切りで並べても構いませんが、**項目を飛ばさないこと**を最優先してください。

【粒度のイメージ（具体例）】
- FAXの例：「この画面でこの項目を確認する」「このタイミングでこの担当者が、黄色のファイルにこの申請書を綴じる」
- システム操作の例：「担当者Aが受付管理画面の当日受付タブを開き、`申請番号`フィールドに申請書の番号を入力して`検索`ボタンを押し、ステータスを確認する」
- 書類処理の例：「窓口担当が受付印を押した申請書を、申請種別ごとに色分けされたファイル（青：新規、黄：変更、赤：取消）に時系列で綴じる」

【重要な指示 - 不足情報の明示】
**画面や音声から、フィールド名や具体的な値が読み取れない場合は、description に曖昧な表現をねじ込まず、ai_hypothesis に明示してください。**

不足情報の書き方の例：
- 「【情報不足】この操作を行う担当者（役職・部署）が不明です。顧客へのヒアリングが必要です。」
- 「【情報不足】`受付日`フィールドにどの値を入れるか（申請書の日付？当日の日付？）が不明です。」
- 「【情報不足】書類を綴じるファイルの色や種類が言及されていません。」
- 「【情報不足】検索条件として使用する項目名が特定できません。顧客に確認が必要です。」
- 「【情報不足】この判断の基準となる金額や条件が明示されていません。例：◯円以上の場合など。」

【入力（全文文字起こし）】
{full_transcript}

【出力フォーマット（必ずこのJSONだけを返すこと）】
{{
  "flow_summary": "業務全体の流れを2〜4文で説明した要約。この業務が何を達成しようとしているか、どんなステップを経るかを簡潔に説明する。",
  "actions": [
    {{
      "index": 1,
      "action_title": "アクションのタイトル（例：受付管理画面で申請書を検索する）",
      "description": "このアクションの目的、具体的な操作、操作後に何が起きるかを含む完結した説明文。`フィールド名`や`ボタン名`を明記し、誰が・どの画面で・何を・どう操作するかを具体的に書く。",
      "actor": "この操作を行う担当者・役職・部署（不明な場合は空文字）",
      "screen": "操作する画面名・システム名・タブ名（不明な場合は空文字）",
      "object": "操作対象となる書類・データ・項目（不明な場合は空文字）",
      "operation": "具体的な操作内容（クリック、入力、確認、分類など）",
      "ai_hypothesis": "【要確認】や【情報不足】として、確認が必要な事項や不明点を質問形式で記載。情報が十分な場合は空文字。",
      "importance": "high" または "normal"（重要な判断ポイント・リスクがある操作なら high）
    }}
  ]
}}

注意:
- 必ずJSON形式のみを返し、余計な文章や説明は一切付けないでください。
- actions配列の各要素にはindexを1から順番に振ってください。
- **タイムスタンプはこの段階では不要です。** アクションの内容だけに集中してください。
- descriptionは「〜します」「〜を確認します」のような丁寧な文体で記述してください。
- **アクションの数が多くなっても構いません。細かく分割することを優先してください。**
- **入力・選択操作では、`フィールド名`と入力する値を具体的に記述してください。**
- **項目名やリストは「など」「等」で省略せず、すべて列挙してください。**
- 情報がある部分は可能な限り具体的に、情報がない部分は「【情報不足】」として見える化してください。
"""

        response = client.models.generate_content(
            model=GEMINI_MODEL_PRO,
            contents=prompt,
        )
        response_text = response.text.strip()

        # コードブロック対策
        if "```json" in response_text:
            response_text = response_text.split("```json")[1].split("```")[0]
        elif "```" in response_text:
            response_text = response_text.split("```")[1].split("```")[0]

        result = json.loads(response_text)
        return result

    except json.JSONDecodeError as e:
        st.warning(f"Gemini応答のJSONパースに失敗しました: {str(e)}")
        return None
    except Exception as e:
        st.error(f"Gemini API エラー: {str(e)}")
        return None


def calculate_text_similarity(text1: str, text2: str) -> float:
    """
    2つのテキスト間の簡易的な類似度を計算（単語の重複率）
    """
    # 日本語用：文字単位でn-gramを作成
    def get_ngrams(text, n=2):
        text = text.lower().strip()
        # 句読点や記号を除去
        text = re.sub(r'[、。！？\s\[\]\-:：]', '', text)
        return set(text[i:i+n] for i in range(len(text) - n + 1)) if len(text) >= n else {text}

    ngrams1 = get_ngrams(text1)
    ngrams2 = get_ngrams(text2)

    if not ngrams1 or not ngrams2:
        return 0.0

    intersection = ngrams1 & ngrams2
    union = ngrams1 | ngrams2

    return len(intersection) / len(union) if union else 0.0


def map_actions_to_segments(actions: list, segments: list) -> list:
    """
    【Stage 2】アクション列をセグメント群にマッピングしてタイムスタンプを付与

    各アクションに対して、最も関連性の高いセグメントを見つけ、
    そのセグメントの時間範囲を使ってタイムスタンプを設定する。

    マッピングロジック:
    1. アクションのdescriptionとaction_titleから重要キーワードを抽出
    2. 各セグメントとの類似度を計算
    3. 最も類似度が高いセグメント群を特定
    4. アクションの順序を保ちながら、時間的に前進するようにマッピング
    """
    if not segments:
        return actions

    mapped_actions = []
    last_used_segment_idx = -1
    total_duration = segments[-1]["end"] if segments else 0

    # 各アクションに均等に時間を割り当てる（フォールバック用）
    time_per_action = total_duration / len(actions) if actions else 0

    for action_idx, action in enumerate(actions):
        action_text = f"{action.get('action_title', '')} {action.get('description', '')}"

        best_segment_idx = -1
        best_score = -1

        # 前のアクションより後のセグメントのみを検索（順序を保つ）
        search_start = max(0, last_used_segment_idx)

        for seg_idx in range(search_start, len(segments)):
            seg = segments[seg_idx]
            seg_text = seg.get("text", "")

            score = calculate_text_similarity(action_text, seg_text)

            # 時間的な位置による重み付け（アクションの順序に近いセグメントを優先）
            expected_time = action_idx * time_per_action
            seg_mid_time = (seg["start"] + seg["end"]) / 2
            time_distance = abs(seg_mid_time - expected_time)
            time_weight = 1.0 / (1.0 + time_distance / total_duration * 2) if total_duration > 0 else 1.0

            weighted_score = score * 0.7 + time_weight * 0.3

            if weighted_score > best_score:
                best_score = weighted_score
                best_segment_idx = seg_idx

        # マッチするセグメントが見つからない場合は、均等割り当てを使用
        if best_segment_idx < 0:
            start_sec = action_idx * time_per_action
            end_sec = (action_idx + 1) * time_per_action
        else:
            # 見つかったセグメントとその周辺を使用
            matched_seg = segments[best_segment_idx]

            # 前後のセグメントも含めて時間範囲を拡張（同じアクションに関連する可能性）
            start_seg_idx = best_segment_idx
            end_seg_idx = best_segment_idx

            # 類似度が一定以上の隣接セグメントを含める
            threshold = best_score * 0.5 if best_score > 0 else 0

            # 前方を確認
            for i in range(best_segment_idx - 1, max(last_used_segment_idx, -1), -1):
                seg_text = segments[i].get("text", "")
                if calculate_text_similarity(action_text, seg_text) >= threshold:
                    start_seg_idx = i
                else:
                    break

            # 後方を確認（次のアクションとの重複を避けるため、控えめに）
            for i in range(best_segment_idx + 1, min(best_segment_idx + 3, len(segments))):
                seg_text = segments[i].get("text", "")
                if calculate_text_similarity(action_text, seg_text) >= threshold:
                    end_seg_idx = i
                else:
                    break

            start_sec = segments[start_seg_idx]["start"]
            end_sec = segments[end_seg_idx]["end"]
            last_used_segment_idx = end_seg_idx

        # マッピング結果をアクションに追加
        mid_sec = (start_sec + end_sec) / 2

        mapped_action = {
            "index": action.get("index", action_idx + 1),
            "action_title": action.get("action_title", f"手順 {action_idx + 1}"),
            "description": action.get("description", ""),
            "actor": action.get("actor", ""),
            "screen": action.get("screen", ""),
            "object": action.get("object", ""),
            "operation": action.get("operation", ""),
            "ai_hypothesis": action.get("ai_hypothesis", ""),
            "importance": action.get("importance", "normal"),
            "is_important": action.get("importance") == "high",
            "start_sec": start_sec,
            "end_sec": end_sec,
            "timestamp_seconds": mid_sec,
            "timestamp": format_timestamp(mid_sec),
            "image": None,  # 後でスクリーンショットを取得
        }

        mapped_actions.append(mapped_action)

    return mapped_actions


def extract_screenshots_for_actions(video_path: str, actions: list) -> list:
    """
    マッピング済みのアクション列に対して、代表スクリーンショットを取得
    """
    for action in actions:
        mid_sec = action.get("timestamp_seconds", 0)
        frame = extract_frame(video_path, mid_sec)
        action["image"] = frame

    return actions


def process_video_chunk(chunk_path: str, start_offset: float, groq_api_key: str,
                        whisper_model: str = "whisper-large-v3",
                        progress_callback=None) -> list:
    """
    動画チャンクを処理してセグメントを返す（Groq Whisper APIで音声認識）
    Returns: segments
    """
    all_segments = []

    # 音声抽出（mp3形式で出力 - Groq APIに対応）
    if progress_callback:
        progress_callback("音声を抽出中...")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp_audio:
        audio_path = tmp_audio.name

    # mp3形式で音声抽出（Groq API対応のため）
    try:
        cmd = [
            FFMPEG_PATH, "-y",
            "-i", chunk_path,
            "-vn",
            "-acodec", "libmp3lame",
            "-ar", "16000",
            "-ac", "1",
            "-b:a", "64k",
            audio_path
        ]
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            st.error(f"音声抽出エラー: {result.stderr}")
            return []
    except Exception as e:
        st.error(f"音声抽出エラー: {str(e)}")
        return []

    # Groq Whisper APIで音声認識
    if progress_callback:
        progress_callback("Groq Whisper APIで音声認識中...")

    segments = transcribe_audio_with_groq(audio_path, groq_api_key, whisper_model)

    # タイムスタンプをオフセット調整
    for segment in segments:
        segment["start"] += start_offset
        segment["end"] += start_offset
        all_segments.append(segment)

    # クリーンアップ
    try:
        os.unlink(audio_path)
    except:
        pass

    return all_segments


def generate_fallback_steps(segments: list, video_path: str) -> list:
    """
    Gemini解析が失敗した場合のフォールバック：セグメント単位で簡易ステップを生成
    """
    steps = []
    for i, seg in enumerate(segments):
        mid_sec = (seg["start"] + seg["end"]) / 2.0
        frame = extract_frame(video_path, mid_sec)
        steps.append({
            "index": i + 1,
            "start_sec": seg["start"],
            "end_sec": seg["end"],
            "timestamp_seconds": mid_sec,
            "timestamp": format_timestamp(mid_sec),
            "action_title": f"手順 {i+1}",
            "description": seg.get("text", "（説明を入力してください）"),
            "actor": "",
            "screen": "",
            "object": "",
            "operation": "",
            "ai_hypothesis": "",
            "importance": "normal",
            "is_important": False,
            "image": frame,
        })
    return steps


def generate_markdown(steps: list, flow_summary: str = None) -> tuple:
    """Markdown形式のマニュアルと画像を生成"""
    markdown_lines = ["# 業務マニュアル\n"]

    # フロー概要があれば追加
    if flow_summary:
        markdown_lines.append("## 概要\n")
        markdown_lines.append(f"{flow_summary}\n")
        markdown_lines.append("\n---\n")

    markdown_lines.append("## 手順\n")

    images = {}
    img_counter = 1

    for i, step in enumerate(steps):
        if step.get("deleted", False):
            continue

        markdown_lines.append(f"### {step.get('index', i+1)}. {step['action_title']}\n")
        markdown_lines.append(f"**タイムスタンプ:** {step['timestamp']}\n")

        if step.get("image") is not None:
            img_filename = f"img_{img_counter}.png"
            images[img_filename] = step["image"]
            markdown_lines.append(f"![手順{step.get('index', i+1)}の画像](./{img_filename})\n")
            img_counter += 1

        # 詳細情報（actor, screen, object, operation）があれば表示
        details = []
        if step.get("actor"):
            details.append(f"**担当:** {step['actor']}")
        if step.get("screen"):
            details.append(f"**画面:** {step['screen']}")
        if step.get("object"):
            details.append(f"**対象:** {step['object']}")
        if step.get("operation"):
            details.append(f"**操作:** {step['operation']}")

        if details:
            markdown_lines.append("\n" + " | ".join(details) + "\n")

        markdown_lines.append(f"\n{step['description']}\n")

        if step.get("ai_hypothesis") and not step.get("hypothesis_resolved", False):
            # 【情報不足】や【要確認】を強調表示
            hypothesis = step['ai_hypothesis']
            if "【情報不足】" in hypothesis:
                markdown_lines.append(f"\n> **[情報不足]** {hypothesis.replace('【情報不足】', '')}\n")
            elif "【要確認】" in hypothesis:
                markdown_lines.append(f"\n> **[要確認]** {hypothesis.replace('【要確認】', '')}\n")
            else:
                markdown_lines.append(f"\n> **AI確認事項:** {hypothesis}\n")

        if step.get("is_important") or step.get("importance") == "high":
            markdown_lines.append("\n**[重要ポイント]**\n")

        markdown_lines.append("\n---\n")

    return "\n".join(markdown_lines), images


def create_download_zip(markdown: str, images: dict) -> bytes:
    """MarkdownとファイルをZIP圧縮"""
    buffer = io.BytesIO()

    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("manual.md", markdown)

        for filename, image in images.items():
            img_buffer = io.BytesIO()
            pil_image = Image.fromarray(image)
            pil_image.save(img_buffer, format="PNG")
            zf.writestr(filename, img_buffer.getvalue())

    buffer.seek(0)
    return buffer.getvalue()


def generate_word_document(steps: list, flow_summary: str = None) -> bytes:
    """Word形式（.docx）のマニュアルを生成"""
    if not DOCX_AVAILABLE:
        return None

    doc = Document()

    # タイトル
    title = doc.add_heading('業務マニュアル', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # フロー概要
    if flow_summary:
        doc.add_heading('概要', level=1)
        doc.add_paragraph(flow_summary)
        doc.add_paragraph()  # 空行

    # 手順
    doc.add_heading('手順', level=1)

    for i, step in enumerate(steps):
        if step.get("deleted", False):
            continue

        # 手順タイトル
        step_index = step.get('index', i+1)
        importance_mark = "[重要] " if step.get("is_important") or step.get("importance") == "high" else ""
        heading = doc.add_heading(f'{importance_mark}{step_index}. {step["action_title"]}', level=2)

        # タイムスタンプ
        ts_para = doc.add_paragraph()
        ts_run = ts_para.add_run(f'タイムスタンプ: {step["timestamp"]}')
        ts_run.bold = True
        ts_run.font.size = Pt(10)
        ts_run.font.color.rgb = RGBColor(100, 100, 100)

        # 画像を挿入
        if step.get("image") is not None:
            try:
                # numpy配列をPIL画像に変換してバイトストリームに保存
                pil_image = Image.fromarray(step["image"])
                img_buffer = io.BytesIO()
                pil_image.save(img_buffer, format="PNG")
                img_buffer.seek(0)

                # Word文書に画像を追加（幅を指定）
                doc.add_picture(img_buffer, width=Inches(5.5))

                # 画像を中央寄せ
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"[画像の挿入に失敗しました: {str(e)}]")

        # 詳細情報（actor, screen, object, operation）
        details = []
        if step.get("actor"):
            details.append(f"担当: {step['actor']}")
        if step.get("screen"):
            details.append(f"画面: {step['screen']}")
        if step.get("object"):
            details.append(f"対象: {step['object']}")
        if step.get("operation"):
            details.append(f"操作: {step['operation']}")

        if details:
            detail_para = doc.add_paragraph()
            for j, detail in enumerate(details):
                if j > 0:
                    detail_para.add_run("  |  ")
                run = detail_para.add_run(detail)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(80, 80, 80)

        # 説明文
        doc.add_paragraph(step['description'])

        # AI仮説・情報不足
        if step.get("ai_hypothesis") and not step.get("hypothesis_resolved", False):
            hypothesis = step['ai_hypothesis']
            hypo_para = doc.add_paragraph()

            if "【情報不足】" in hypothesis:
                run = hypo_para.add_run("[情報不足] ")
                run.bold = True
                run.font.color.rgb = RGBColor(200, 0, 0)  # 赤色
                hypo_para.add_run(hypothesis.replace('【情報不足】', ''))
            elif "【要確認】" in hypothesis:
                run = hypo_para.add_run("[要確認] ")
                run.bold = True
                run.font.color.rgb = RGBColor(200, 150, 0)  # オレンジ色
                hypo_para.add_run(hypothesis.replace('【要確認】', ''))
            else:
                run = hypo_para.add_run("AI確認事項: ")
                run.bold = True
                hypo_para.add_run(hypothesis)

        # 重要ポイントマーク
        if step.get("is_important") or step.get("importance") == "high":
            important_para = doc.add_paragraph()
            run = important_para.add_run("[重要ポイント]")
            run.bold = True
            run.font.color.rgb = RGBColor(200, 0, 0)

        # 区切り線代わりの空行
        doc.add_paragraph()
        doc.add_paragraph("─" * 50)
        doc.add_paragraph()

    # バイトストリームとして出力
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def main():
    st.set_page_config(
        page_title="業務マニュアル自動生成",
        page_icon="📋",
        layout="wide"
    )

    # Session State初期化
    if "steps" not in st.session_state:
        st.session_state.steps = []
    if "flow_summary" not in st.session_state:
        st.session_state.flow_summary = None
    if "video_path" not in st.session_state:
        st.session_state.video_path = None
    if "video_duration" not in st.session_state:
        st.session_state.video_duration = 0
    if "segments" not in st.session_state:
        st.session_state.segments = []
    if "processing" not in st.session_state:
        st.session_state.processing = False
    if "selected_step" not in st.session_state:
        st.session_state.selected_step = 0
    if "current_project_id" not in st.session_state:
        st.session_state.current_project_id = None
    if "current_project_name" not in st.session_state:
        st.session_state.current_project_name = ""
    if "video_filename" not in st.session_state:
        st.session_state.video_filename = ""
    if "is_audio_only" not in st.session_state:
        st.session_state.is_audio_only = False
    if "groq_api_key" not in st.session_state:
        st.session_state.groq_api_key = ""
    if "gemini_api_key" not in st.session_state:
        st.session_state.gemini_api_key = ""
    if "whisper_model" not in st.session_state:
        st.session_state.whisper_model = "whisper-large-v3"
    if "supabase_url" not in st.session_state:
        st.session_state.supabase_url = ""
    if "supabase_key" not in st.session_state:
        st.session_state.supabase_key = ""

    # サイドバー
    with st.sidebar:
        # ============================================
        # プロジェクト管理セクション
        # ============================================
        st.header("プロジェクト")

        # プロジェクト一覧を取得
        projects = get_project_list()

        # 新規プロジェクト作成ボタン
        if st.button("➕ 新規プロジェクト", use_container_width=True):
            st.session_state.current_project_id = None
            st.session_state.current_project_name = ""
            st.session_state.steps = []
            st.session_state.flow_summary = None
            st.session_state.video_path = None
            st.session_state.video_duration = 0
            st.session_state.segments = []
            st.session_state.video_filename = ""
            st.session_state.is_audio_only = False
            st.rerun()

        # プロジェクト一覧表示
        if projects:
            st.caption(f"保存済み: {len(projects)}件")

            for proj in projects:
                col1, col2 = st.columns([4, 1])
                with col1:
                    # プロジェクト選択ボタン
                    is_current = (st.session_state.current_project_id == proj["id"])
                    button_label = f"{'📂 ' if is_current else '📁 '}{proj['name'][:20]}"
                    if st.button(button_label, key=f"proj_{proj['id']}", use_container_width=True):
                        # プロジェクトを読み込み
                        data = load_project(proj["id"])
                        if data:
                            st.session_state.current_project_id = proj["id"]
                            st.session_state.current_project_name = data.get("name", "")
                            st.session_state.steps = data.get("steps", [])
                            st.session_state.flow_summary = data.get("flow_summary")
                            st.session_state.segments = data.get("segments", [])
                            st.session_state.video_duration = data.get("video_duration", 0)
                            st.session_state.video_filename = data.get("video_name", "")
                            st.session_state.video_path = data.get("video_path")
                            st.session_state.is_audio_only = data.get("is_audio_only", False)
                            st.rerun()
                with col2:
                    # 削除ボタン
                    if st.button("🗑", key=f"del_{proj['id']}", help="削除"):
                        delete_project(proj["id"])
                        if st.session_state.current_project_id == proj["id"]:
                            st.session_state.current_project_id = None
                            st.session_state.current_project_name = ""
                            st.session_state.steps = []
                            st.session_state.flow_summary = None
                        st.rerun()
        else:
            st.caption("保存されたプロジェクトはありません")

        st.divider()

        # ============================================
        # 設定セクション
        # ============================================
        st.header("API設定")

        # secrets.tomlのパス
        secrets_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".streamlit", "secrets.toml")

        # 保存済みの値を読み込み
        def load_saved_secrets():
            saved = {}
            if os.path.exists(secrets_path):
                try:
                    with open(secrets_path, "r") as f:
                        for line in f:
                            line = line.strip()
                            if "=" in line and not line.startswith("#"):
                                key, val = line.split("=", 1)
                                saved[key.strip()] = val.strip().strip('"')
                except:
                    pass
            return saved

        saved_secrets = load_saved_secrets()
        default_groq_key = saved_secrets.get("GROQ_API_KEY", "")
        default_gemini_key = saved_secrets.get("GEMINI_API_KEY", "")
        default_supabase_url = saved_secrets.get("SUPABASE_URL", "")
        default_supabase_key = saved_secrets.get("SUPABASE_KEY", "")

        groq_api_key = st.text_input(
            "Groq API Key",
            value=default_groq_key,
            type="password",
            help="https://console.groq.com/ でAPIキーを取得（無料）"
        )

        gemini_api_key = st.text_input(
            "Gemini API Key",
            value=default_gemini_key,
            type="password",
            help="Google AI StudioでAPIキーを取得してください"
        )

        with st.expander("クラウド保存設定（Supabase）"):
            supabase_url = st.text_input(
                "Supabase URL",
                value=default_supabase_url,
                help="Supabaseプロジェクトの URL"
            )
            supabase_key = st.text_input(
                "Supabase Anon Key",
                value=default_supabase_key,
                type="password",
                help="Supabaseの anon/public キー"
            )

            # session_stateに保存
            if supabase_url:
                st.session_state.supabase_url = supabase_url
            if supabase_key:
                st.session_state.supabase_key = supabase_key

            if supabase_url and supabase_key:
                st.success("Supabase接続済み")
            else:
                st.info("Supabaseを設定するとプロジェクトがクラウドに保存されます")

        # APIキー保存ボタン（ローカル環境のみ有効）
        is_cloud = os.environ.get("STREAMLIT_SHARING_MODE") or os.path.exists("/mount/src")
        if is_cloud:
            st.info("クラウド環境ではStreamlit CloudのSecretsで管理されています")
        else:
            if st.button("APIキーを保存", use_container_width=True):
                try:
                    os.makedirs(os.path.dirname(secrets_path), exist_ok=True)
                    with open(secrets_path, "w") as f:
                        if groq_api_key:
                            f.write(f'GROQ_API_KEY = "{groq_api_key}"\n')
                        if gemini_api_key:
                            f.write(f'GEMINI_API_KEY = "{gemini_api_key}"\n')
                        if supabase_url:
                            f.write(f'SUPABASE_URL = "{supabase_url}"\n')
                        if supabase_key:
                            f.write(f'SUPABASE_KEY = "{supabase_key}"\n')
                    st.success("保存しました！次回から自動入力されます")
                except Exception as e:
                    st.error(f"保存に失敗: {str(e)}")

        whisper_model = st.selectbox(
            "Whisperモデル",
            options=["whisper-large-v3", "whisper-large-v3-turbo"],
            index=0,
            help="large-v3: 最高精度 / large-v3-turbo: 高速・低コスト",
            format_func=lambda x: "Whisper large-v3（最高精度）" if x == "whisper-large-v3" else "Whisper large-v3 Turbo（高速）"
        )

        st.divider()

        st.caption(f"音声認識: Groq {whisper_model}")
        st.caption(f"テキスト分析: {GEMINI_MODEL_PRO}")

        st.divider()

        # ============================================
        # 動画アップロードセクション
        # ============================================
        st.header("ファイルアップロード")

        uploaded_file = st.file_uploader(
            "ファイルをアップロード（MP4/MP3）",
            key="file_uploader_v2",
            help="MP4（動画）またはMP3（音声）ファイルをアップロードしてください（大きなファイルは自動分割されます）"
        )

        if uploaded_file is not None:
            # ファイル形式チェック
            file_ext = os.path.splitext(uploaded_file.name)[1].lower()
            if file_ext not in [".mp4", ".mp3"]:
                st.error("MP4またはMP3ファイルをアップロードしてください")
                uploaded_file = None

        if uploaded_file is not None:
            # ファイルサイズ表示
            file_size_mb = uploaded_file.size / (1024 * 1024)
            st.info(f"ファイルサイズ: {file_size_mb:.1f} MB")

            if file_size_mb > MAX_CHUNK_SIZE_MB:
                st.warning(f"大きなファイルです。{MAX_CHUNK_SIZE_MB}MB以下に自動分割して処理します。")

            # 一時ファイルとして保存
            is_audio = uploaded_file.name.lower().endswith(".mp3")
            file_suffix = ".mp3" if is_audio else ".mp4"
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_suffix) as tmp:
                tmp.write(uploaded_file.read())
                st.session_state.video_path = tmp.name
                st.session_state.is_audio_only = is_audio
                if is_audio:
                    st.session_state.video_duration = get_audio_duration(tmp.name)
                else:
                    st.session_state.video_duration = get_video_duration(tmp.name)
                st.session_state.video_filename = uploaded_file.name

            if st.button("解析開始", type="primary", use_container_width=True):
                if not groq_api_key:
                    st.error("Groq API Keyを入力してください")
                elif not gemini_api_key:
                    st.error("Gemini API Keyを入力してください")
                else:
                    # 新しいプロジェクトIDを生成
                    st.session_state.current_project_id = generate_project_id()
                    st.session_state.current_project_name = os.path.splitext(uploaded_file.name)[0]
                    st.session_state.processing = True
                    st.session_state.groq_api_key = groq_api_key
                    st.session_state.gemini_api_key = gemini_api_key
                    st.session_state.whisper_model = whisper_model

                    # 動画/音声をprojects/に永続化コピー
                    ext = ".mp3" if is_audio else ".mp4"
                    dest = _project_source_path(st.session_state.current_project_id, ext)
                    shutil.copy2(st.session_state.video_path, dest)
                    st.session_state.video_path = dest

    # メインエリアのタイトル
    if st.session_state.current_project_name:
        st.title(f"📋 {st.session_state.current_project_name}")
    else:
        st.title("業務マニュアル自動生成アプリ")

    # メインコンテンツ
    if st.session_state.processing and st.session_state.video_path:
        with st.spinner("ファイルを解析中..."):
            progress = st.progress(0)
            status = st.empty()

            # 【Stage 1-A】音声認識（MP3は直接処理、MP4は音声抽出後に処理）
            if st.session_state.is_audio_only:
                all_segments = transcribe_audio_file(
                    audio_path=st.session_state.video_path,
                    groq_api_key=st.session_state.groq_api_key,
                    whisper_model=st.session_state.whisper_model,
                    progress_callback=lambda p: progress.progress(p),
                    status_callback=lambda s: status.text(s)
                )
            else:
                all_segments = transcribe_full_video(
                    video_path=st.session_state.video_path,
                    groq_api_key=st.session_state.groq_api_key,
                    whisper_model=st.session_state.whisper_model,
                    progress_callback=lambda p: progress.progress(p),
                    status_callback=lambda s: status.text(s)
                )

            st.session_state.segments = all_segments

            # 【Stage 1-B】フルコンテクストをGemini Proに渡してアクション列を生成
            status.text("AIが業務フローを分析中...")
            progress.progress(45)

            full_transcript = build_full_transcript(all_segments)

            if not full_transcript.strip():
                st.warning("音声が認識されませんでした。ファイルに音声が含まれているか確認してください。")
                st.session_state.steps = []
                st.session_state.flow_summary = None
                st.session_state.processing = False
                st.rerun()

            # Gemini Pro でフルコンテクスト解析
            print(f"[DEBUG] 文字起こしセグメント数: {len(all_segments)}")
            print(f"[DEBUG] 文字起こし文字数: {len(full_transcript)}")
            gemini_result = analyze_full_context_with_gemini(st.session_state.gemini_api_key, full_transcript)

            if gemini_result is None:
                # フォールバック：セグメント単位で簡易ステップを生成
                status.text("フォールバック処理中...")
                progress.progress(70)
                print("[DEBUG] Gemini解析失敗 → フォールバック処理")

                st.warning("AI解析に失敗したため、簡易な手順を生成します。")
                steps = generate_fallback_steps(all_segments, st.session_state.video_path)
                st.session_state.flow_summary = None
            else:
                flow_summary = gemini_result.get("flow_summary", "")
                actions = gemini_result.get("actions", [])
                print(f"[DEBUG] Gemini解析成功: flow_summary長={len(flow_summary)}, actions数={len(actions)}")

                st.session_state.flow_summary = flow_summary

                # 【Stage 2-A】アクションをセグメントにマッピング
                status.text("アクションと映像をマッピング中...")
                progress.progress(60)

                mapped_actions = map_actions_to_segments(actions, all_segments)

                # 【Stage 2-B】スクリーンショットを取得（動画の場合のみ）
                if not st.session_state.is_audio_only:
                    status.text("スクリーンショットを取得中...")
                    progress.progress(80)
                    steps = extract_screenshots_for_actions(st.session_state.video_path, mapped_actions)
                else:
                    progress.progress(80)
                    steps = mapped_actions

            st.session_state.steps = steps

            progress.progress(95)
            status.text("プロジェクトを保存中...")

            # プロジェクトを自動保存
            if st.session_state.current_project_id:
                save_project(
                    project_id=st.session_state.current_project_id,
                    name=st.session_state.current_project_name,
                    video_name=st.session_state.video_filename,
                    steps=st.session_state.steps,
                    flow_summary=st.session_state.flow_summary,
                    segments=st.session_state.segments,
                    video_duration=st.session_state.video_duration
                )

            progress.progress(100)
            status.text("完了!")

            st.session_state.processing = False
            st.rerun()

    # エディタUI
    if st.session_state.steps:
        # プロジェクト名編集と保存ボタン
        proj_col1, proj_col2, proj_col3 = st.columns([3, 1, 1])
        with proj_col1:
            new_project_name = st.text_input(
                "プロジェクト名",
                value=st.session_state.current_project_name,
                key="project_name_input",
                label_visibility="collapsed",
                placeholder="プロジェクト名を入力..."
            )
            if new_project_name != st.session_state.current_project_name:
                st.session_state.current_project_name = new_project_name
        with proj_col2:
            if st.button("💾 保存", use_container_width=True, type="primary"):
                if st.session_state.current_project_id:
                    save_project(
                        project_id=st.session_state.current_project_id,
                        name=st.session_state.current_project_name,
                        video_name=st.session_state.video_filename,
                        steps=st.session_state.steps,
                        flow_summary=st.session_state.flow_summary,
                        segments=st.session_state.segments,
                        video_duration=st.session_state.video_duration
                    )
                    st.success("保存しました！")
                else:
                    # プロジェクトIDがない場合は新規作成
                    st.session_state.current_project_id = generate_project_id()
                    save_project(
                        project_id=st.session_state.current_project_id,
                        name=st.session_state.current_project_name or "無題のプロジェクト",
                        video_name=st.session_state.video_filename,
                        steps=st.session_state.steps,
                        flow_summary=st.session_state.flow_summary,
                        segments=st.session_state.segments,
                        video_duration=st.session_state.video_duration
                    )
                    st.success("新規保存しました！")
                    st.rerun()
        with proj_col3:
            st.caption(f"ID: {st.session_state.current_project_id or '未保存'}")

        # フロー概要を表示
        if st.session_state.flow_summary:
            st.info(f"**業務フロー概要:** {st.session_state.flow_summary}")

        col_left, col_right = st.columns([1.5, 1])

        # 右カラム: プレビュー & ツール
        with col_right:
            if not st.session_state.get("is_audio_only", False):
                st.subheader("動画プレビュー")

                if st.session_state.video_path and os.path.exists(st.session_state.video_path):
                    st.video(st.session_state.video_path)

                st.divider()

                st.subheader("画像取得ツール")

                capture_time = st.slider(
                    "タイムスタンプ（秒）",
                    min_value=0.0,
                    max_value=float(max(1, int(st.session_state.video_duration))),
                    value=0.0,
                    step=0.5,
                    key="capture_slider"
                )

                target_step = st.selectbox(
                    "適用先の手順",
                    options=range(len(st.session_state.steps)),
                    format_func=lambda x: f"{st.session_state.steps[x].get('index', x+1)}. {st.session_state.steps[x]['action_title']}" if not st.session_state.steps[x].get("deleted") else f"{x+1}. (削除済み)",
                    key="target_step_select"
                )

                if st.button("この瞬間の画像をキャプチャ", use_container_width=True):
                    if st.session_state.video_path:
                        new_frame = extract_frame(st.session_state.video_path, capture_time)
                        if new_frame is not None:
                            st.session_state.steps[target_step]["image"] = new_frame
                            st.session_state.steps[target_step]["timestamp"] = format_timestamp(capture_time)
                            st.session_state.steps[target_step]["timestamp_seconds"] = capture_time
                            st.success("画像を更新しました")
                            st.rerun()
                        else:
                            st.error("フレームの取得に失敗しました")
            else:
                st.subheader("音声プレビュー")
                if st.session_state.video_path and os.path.exists(st.session_state.video_path):
                    st.audio(st.session_state.video_path)
                st.info("音声ファイルのためスクリーンショット取得は利用できません")

        # 左カラム: マニュアルエディタ
        with col_left:
            st.subheader("マニュアル編集")

            for i, step in enumerate(st.session_state.steps):
                if step.get("deleted", False):
                    continue

                with st.container():
                    importance_badge = "🔴 " if step.get("is_important") or step.get("importance") == "high" else ""
                    st.markdown(f"### {importance_badge}手順 {step.get('index', i+1)}: {step['timestamp']}")

                    col_img, col_edit = st.columns([1, 2])

                    with col_img:
                        if step.get("image") is not None:
                            st.image(step["image"], use_container_width=True)
                            if st.button("画像を削除", key=f"del_img_{i}"):
                                st.session_state.steps[i]["image"] = None
                                st.rerun()
                        else:
                            st.info("画像なし")

                    with col_edit:
                        new_title = st.text_input(
                            "タイトル",
                            value=step["action_title"],
                            key=f"title_{i}"
                        )
                        st.session_state.steps[i]["action_title"] = new_title

                        # 詳細情報（actor, screen, object, operation）を表示
                        detail_cols = st.columns(2)
                        with detail_cols[0]:
                            if step.get("actor"):
                                st.caption(f"担当: {step['actor']}")
                            if step.get("screen"):
                                st.caption(f"画面: {step['screen']}")
                        with detail_cols[1]:
                            if step.get("object"):
                                st.caption(f"対象: {step['object']}")
                            if step.get("operation"):
                                st.caption(f"操作: {step['operation']}")

                        new_desc = st.text_area(
                            "説明",
                            value=step["description"],
                            key=f"desc_{i}",
                            height=100
                        )
                        st.session_state.steps[i]["description"] = new_desc

                        if step.get("ai_hypothesis"):
                            hypothesis = step['ai_hypothesis']
                            # 【情報不足】や【要確認】を強調表示
                            if "【情報不足】" in hypothesis:
                                st.error(f"情報不足: {hypothesis}")
                            elif "【要確認】" in hypothesis:
                                st.warning(f"要確認: {hypothesis}")
                            elif step.get("is_important") or step.get("importance") == "high":
                                st.warning(f"AI仮説: {hypothesis}")
                            else:
                                st.info(f"AI仮説: {hypothesis}")

                            resolved = st.checkbox(
                                "解決済み",
                                value=step.get("hypothesis_resolved", False),
                                key=f"resolved_{i}"
                            )
                            st.session_state.steps[i]["hypothesis_resolved"] = resolved

                        if st.button("この手順を削除", key=f"del_step_{i}"):
                            st.session_state.steps[i]["deleted"] = True
                            st.rerun()

                    st.divider()

            # ダウンロード
            st.subheader("出力")

            # 出力形式の選択
            output_format = st.radio(
                "出力形式",
                options=["Word (.docx)", "Markdown (.md)"],
                horizontal=True,
                key="output_format"
            )

            if output_format == "Word (.docx)":
                if DOCX_AVAILABLE:
                    if st.button("Wordファイルを生成", type="primary", use_container_width=True):
                        word_data = generate_word_document(
                            st.session_state.steps,
                            st.session_state.flow_summary
                        )
                        if word_data:
                            st.download_button(
                                label="Wordファイルをダウンロード (.docx)",
                                data=word_data,
                                file_name="manual.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                        else:
                            st.error("Word文書の生成に失敗しました")
                else:
                    st.warning("python-docxがインストールされていません。`pip install python-docx`を実行してください。")
            else:
                if st.button("Markdownを生成", type="primary", use_container_width=True):
                    markdown, images = generate_markdown(
                        st.session_state.steps,
                        st.session_state.flow_summary
                    )
                    zip_data = create_download_zip(markdown, images)

                    st.download_button(
                        label="ZIPファイルをダウンロード (.md + 画像)",
                        data=zip_data,
                        file_name="manual.zip",
                        mime="application/zip",
                        use_container_width=True
                    )

    elif not st.session_state.processing:
        st.info("サイドバーから動画をアップロードして解析を開始してください。")

        with st.expander("使い方"):
            st.markdown("""
            ### 基本的な使い方
            1. **サイドバー**でGemini API Keyを入力
            2. 必要に応じてWhisperモデルを調整
            3. MP4動画またはMP3音声をアップロード
            4. 「解析開始」ボタンをクリック
            5. 生成されたマニュアルを編集
            6. Word または Markdown でダウンロード

            ### プロジェクト管理
            - **自動保存**: 解析完了後、プロジェクトは自動的に保存されます
            - **手動保存**: 編集後は「💾 保存」ボタンで保存できます
            - **プロジェクト切り替え**: サイドバーの📁アイコンをクリックで切り替え
            - **プロジェクト削除**: 🗑アイコンで削除
            - **新規作成**: 「➕ 新規プロジェクト」ボタンで新しいプロジェクトを開始

            ### 処理方式
            - **Stage 1**: 動画全体の文字起こしをAI（Gemini Pro）に渡し、業務フロー全体を理解させます
            - **Stage 2**: AIが生成したアクション列に対して、適切なスクリーンショットを自動マッピングします

            これにより、従来の「ぶつ切り」な説明ではなく、**文脈を踏まえた完結した説明文**が生成されます。

            ### 大きなファイルについて
            100MBを超えるファイルは自動的に分割して処理され、結果が統合されます。
            """)


if __name__ == "__main__":
    main()
